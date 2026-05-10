"""Pass 1 — deterministic extraction from .docx and .xlsx (no API calls)."""

from __future__ import annotations

import re
import warnings
from pathlib import Path
from typing import Any

import openpyxl
from docx import Document
from openpyxl.chartsheet.chartsheet import Chartsheet

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ---------------------------------------------------------------------------
# DOCX extraction — helpers
# ---------------------------------------------------------------------------

def _para_style(para) -> str:
    return para.style.name if para.style else ""


def _is_heading(para) -> bool:
    return _para_style(para).startswith("Heading")


def _heading_level(para) -> int:
    m = re.match(r"Heading (\d+)", _para_style(para))
    return int(m.group(1)) if m else 0


def _para_text(para) -> str:
    return para.text.strip()


# ---------------------------------------------------------------------------
# TOC extraction — three strategies, in priority order
# ---------------------------------------------------------------------------

def _extract_toc_from_sdt(doc: Document) -> dict[str, int]:
    """Extract the section TOC from a Structured Document Tag (SDT).

    Word stores the auto-generated section TOC in an SDT element whose content
    is a flat list of text runs grouped as triples: section_no, title, page.
    E.g. '1', 'EXECUTIVE SUMMARY', '12' → section '1 EXECUTIVE SUMMARY' → page 12.

    Returns {full_section_title: page_number}.
    """
    toc: dict[str, int] = {}
    body = doc.element.body
    sdts = body.findall(".//{%s}sdt" % _W_NS)
    for sdt in sdts:
        t_elements = sdt.findall(".//{%s}t" % _W_NS)
        texts = [el.text for el in t_elements if el.text and el.text.strip()]
        if len(texts) < 3:
            continue
        # Heuristic: SDT is a section TOC when ≥30% of runs are page-number-like
        # (short strings that are purely numeric) and the rest are section identifiers.
        numeric = [t for t in texts if re.match(r"^\d+$", t.strip())]
        if len(numeric) < len(texts) * 0.25:
            continue
        # Parse triples: section_no, title, page
        i = 0
        while i < len(texts) - 2:
            sec_no = texts[i].strip()
            title = texts[i + 1].strip()
            page_str = texts[i + 2].strip()
            if re.match(r"^[\d.]+$", sec_no) and re.match(r"^\d+$", page_str):
                full_title = f"{sec_no} {title}"
                toc[full_title] = int(page_str)
                toc[title] = int(page_str)  # also index by bare title
                i += 3
            else:
                i += 1
    return toc


def _extract_toc_from_paragraphs(doc: Document) -> dict[str, int]:
    """Parse TOC entries from paragraph styles (TOC 1/2/3 or 'table of figures').

    These give page numbers for listed tables, figures, and attachments —
    useful supplementary references even when section-level TOC is missing here.
    Returns {caption_or_title: page_number}.
    """
    toc: dict[str, int] = {}
    for para in doc.paragraphs:
        style = _para_style(para)
        is_toc_style = style.startswith("TOC") or style.startswith("toc")
        is_caption_list = "table of figures" in style.lower()
        if not (is_toc_style or is_caption_list):
            continue
        text = _para_text(para)
        if not text:
            continue
        m = re.search(r"\b(\d+)\s*$", text)
        if m:
            page = int(m.group(1))
            title = re.sub(r"[\s.]+\d+\s*$", "", text).strip()
            toc[title] = page
    return toc


def _extract_headings(doc: Document) -> list[dict[str, Any]]:
    """Extract all headings with their level and text."""
    headings = []
    for para in doc.paragraphs:
        if _is_heading(para):
            headings.append({
                "level": _heading_level(para),
                "text": _para_text(para),
            })
    return headings


def _build_section_page_map(
    sdt_toc: dict[str, int],
    caption_toc: dict[str, int],
    headings: list[dict[str, Any]],
) -> dict[str, int | None]:
    """Merge SDT section TOC + caption-list page refs with heading structure.

    Priority: SDT exact → SDT bare title → caption-list → None.
    Returns {heading_text: page | None}.
    """
    # Combined lookup: SDT has priority
    combined = {**caption_toc, **sdt_toc}

    section_map: dict[str, int | None] = {}
    for h in headings:
        title = h["text"]
        if title in combined:
            section_map[title] = combined[title]
            continue
        # Strip leading section number
        stripped = re.sub(r"^\d[\d.]*\s+", "", title).strip()
        if stripped in combined:
            section_map[title] = combined[stripped]
            continue
        # Substring fallback
        matched = None
        for key, page in combined.items():
            if stripped and stripped.lower() in key.lower():
                matched = page
                break
            if title.lower() in key.lower():
                matched = page
                break
        section_map[title] = matched
    return section_map


def _doc_body_text(doc: Document) -> str:
    """Concatenate all paragraph text preserving heading markers."""
    lines: list[str] = []
    for para in doc.paragraphs:
        text = _para_text(para)
        if not text:
            lines.append("")
            continue
        if _is_heading(para):
            prefix = "#" * _heading_level(para)
            lines.append(f"{prefix} {text}")
        else:
            lines.append(text)
    # Also pull text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Public DOCX entry point
# ---------------------------------------------------------------------------

def extract_docx(path: str | Path) -> dict[str, Any]:
    """Extract text, TOC, and section-page map from a .docx report.

    Returns:
        {
          "text": str,                    # full body text (markdown-like)
          "toc": {title: page},           # section TOC from SDT (primary)
          "caption_toc": {title: page},   # table/figure list page refs
          "section_page_map": {title: page | None},
          "headings": [{level, text}, ...],
          "toc_source": "sdt" | "paragraphs" | "none",
        }
    """
    doc = Document(str(path))
    sdt_toc = _extract_toc_from_sdt(doc)
    caption_toc = _extract_toc_from_paragraphs(doc)
    headings = _extract_headings(doc)
    section_page_map = _build_section_page_map(sdt_toc, caption_toc, headings)
    text = _doc_body_text(doc)

    toc_source = "sdt" if sdt_toc else ("paragraphs" if caption_toc else "none")

    return {
        "text": text,
        "toc": sdt_toc,
        "caption_toc": caption_toc,
        "section_page_map": section_page_map,
        "headings": headings,
        "toc_source": toc_source,
    }


# ---------------------------------------------------------------------------
# XLSX extraction
# ---------------------------------------------------------------------------

def _sheet_to_rows(ws) -> list[list[Any]]:
    """Convert worksheet to list of rows (list of cell values)."""
    rows = []
    for row in ws.iter_rows(values_only=True):
        rows.append(list(row))
    return rows


def extract_xlsx(path: str | Path) -> dict[str, Any]:
    """Extract every sheet from an .xlsx workbook.

    Returns:
        {
          "sheet_names": [str, ...],
          "sheets": {sheet_name: [[cell_value, ...], ...]},
          "chart_sheets": [str, ...],
        }
    """
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        wb = openpyxl.load_workbook(str(path), data_only=True)

    sheets: dict[str, list[list[Any]]] = {}
    chart_sheets: list[str] = []

    for name in wb.sheetnames:
        ws = wb[name]
        if isinstance(ws, Chartsheet):
            chart_sheets.append(name)
        else:
            sheets[name] = _sheet_to_rows(ws)

    return {
        "sheet_names": wb.sheetnames,
        "sheets": sheets,
        "chart_sheets": chart_sheets,
    }

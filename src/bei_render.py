"""Render BEI report as Word .docx matching Pavilion template style."""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

_NAVY = RGBColor(0x1A, 0x2E, 0x4A)


def _set_cell_bg(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _bold_cell(cell) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True


def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = _NAVY
    return p


def _para(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    return p


def _star_rating(bei: float) -> str:
    if bei <= 100:  return "5-Star"
    if bei <= 135:  return "4-Star"
    if bei <= 175:  return "3-Star"
    if bei <= 220:  return "2-Star"
    return "1-Star"


def _table2(doc: Document, rows_data: list[tuple[str, str]], style: str = "Table Grid") -> None:
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = style
    for i, (label, value) in enumerate(rows_data):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = str(value)
        _bold_cell(tbl.rows[i].cells[0])


def render_bei_docx(
    profile: dict[str, Any],
    energy_data: dict[str, Any],
    narratives: dict[str, str],
) -> bytes:
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    gfa     = float(profile.get("gfa", 1) or 1)
    total_a = float(energy_data.get("total_a", 0) or 0)
    total_b = float(energy_data.get("total_b", 0) or 0)
    total_c = float(energy_data.get("total_c", 0) or 0)
    bei     = round(total_a / gfa, 2)          # BEI = Total Supply (A) ÷ GFA
    bei_gj  = round(bei * 0.0036, 3)
    months  = energy_data.get("months", [])
    period  = energy_data.get("period_label", "")
    star    = _star_rating(bei)

    building  = profile.get("building_name", "")
    client    = profile.get("client_name", "")
    address   = profile.get("address", "")
    prep_name = profile.get("preparer_name", "Atech Energy Sdn Bhd")
    prep_pos  = profile.get("preparer_position", "Energy Auditor")
    sub_date  = profile.get("submission_date", "")
    yr        = int(profile.get("year_completed") or 2000)
    age       = 2025 - yr
    tariff    = float(profile.get("tariff_rate_sen", 36.5) or 36.5) / 100

    monthly_a = energy_data.get("monthly_a", [])
    monthly_b = energy_data.get("monthly_b", [])
    monthly_c = energy_data.get("monthly_c", [])

    # ── Cover ─────────────────────────────────────────────────────────────────
    _table2(doc, [
        ("Client Name",        client),
        ("Building Name",      building),
        ("Address",            address),
        ("Prepared By",        prep_name),
        ("Date of Submission", sub_date),
    ])
    doc.add_page_break()

    # ── 1. Declaration ────────────────────────────────────────────────────────
    _heading(doc, "1. Declaration")
    _para(doc, (
        f"I hereby declare that the information contained in this Building Energy Intensity (BEI) "
        f"Report for {building} is accurate and complete to the best of my knowledge. "
        f"This report has been prepared in compliance with the applicable energy efficiency "
        f"reporting requirements under the Energy Efficiency and Conservation Act 2024 (EECA 2024) "
        f"and in accordance with the guidelines issued by Suruhanjaya Tenaga (ST) Malaysia."
    ))
    doc.add_paragraph()
    _table2(doc, [
        ("Name",           prep_name),
        ("Position",       prep_pos),
        ("Company",        "Atech Energy Sdn Bhd"),
        ("Date",           sub_date),
        ("Signature",      ""),
        ("Qualifications", ""),
    ])
    doc.add_page_break()

    # ── 2. Executive Summary ──────────────────────────────────────────────────
    _heading(doc, "2. Executive Summary")
    for blk in narratives.get("executive_summary", "").split("\n\n"):
        if blk.strip():
            _para(doc, blk.strip())
    doc.add_page_break()

    # ── 3. Introduction ───────────────────────────────────────────────────────
    _heading(doc, "3. Introduction")
    for blk in narratives.get("intro_narrative", "").split("\n\n"):
        if blk.strip():
            _para(doc, blk.strip())
    doc.add_paragraph()

    _heading(doc, "3.1 Building Information", level=2)
    ac_pct           = float(profile.get("ac_pct", 0) or 0)
    server_area_pct  = float(profile.get("server_area_pct", 0) or 0)
    parking_area_pct = float(profile.get("parking_area_pct", 0) or 0)
    nfa              = float(profile.get("nfa", 0) or 0)
    design_load      = float(profile.get("design_load", 0) or 0)
    design_load_unit = profile.get("design_load_unit", "pax")
    actual_load_pct  = float(profile.get("actual_load_pct", 0) or 0)
    _table2(doc, [
        ("Building Name",                building),
        ("Address",                      address),
        ("Type of Building",             profile.get("building_type", "")),
        ("Year of Completion",           str(yr)),
        ("Age of Building",              f"{age} years"),
        ("Gross Floor Area (GFA)",       f"{gfa:,.2f} m²"),
        ("% of GFA Air Conditioned",     f"{ac_pct:.1f}%"),
        ("Server Area",                  f"{server_area_pct:.1f}%"),
        ("Parking Area (Enclosed)",      f"{parking_area_pct:.1f}%"),
        ("Net Floor Area (NFA)",         f"{nfa:,.2f} m²"),
        ("Design Occupant Load",         f"{design_load:,.0f} {design_load_unit}"),
        ("Actual Occupant Load",         f"{actual_load_pct:.1f}%"),
        ("Certifications",               profile.get("certifications", "None")),
        ("Energy Source",                energy_data.get("supply_auth", "TNB")),
        ("TNB Account No.",              energy_data.get("tnb_account", "")),
        ("Total Supply — Annual (A)",    f"{total_a:,.2f} kWh"),
        ("Net Landlord Load (C = A−B)", f"{total_c:,.2f} kWh"),
        ("Tenant Load (B)",              f"{total_b:,.2f} kWh"),
        ("BEI  (= A ÷ GFA)",            f"{bei} kWh/m²/year  =  {bei_gj} GJ/m²/year"),
        ("ST Star Rating",               star),
    ])
    doc.add_paragraph()

    _heading(doc, "3.2 Building Operating Hours", level=2)
    _table2(doc, [
        ("Operating Hours",        profile.get("operating_hours", "")),
        ("Data Collection Period", period),
    ])
    doc.add_page_break()

    # ── 4. Energy Consumption ─────────────────────────────────────────────────
    _heading(doc, "4. Energy Consumption Data")
    _para(doc, (
        f"The energy consumption data for {building} was collected from {period}. "
        f"Data was obtained from TNB monthly billing statements "
        f"(Account No. {energy_data.get('tnb_account', '')})."
    ))
    doc.add_paragraph()

    _heading(doc, "4.1 Monthly Electricity Consumption (kWh)", level=2)
    n = len(months)
    ec = doc.add_table(rows=n + 2, cols=4)
    ec.style = "Table Grid"
    for j, h in enumerate(["Month", "Total Supply (kWh)", "Landlord (kWh)", "Tenant (kWh)"]):
        ec.rows[0].cells[j].text = h
        _bold_cell(ec.rows[0].cells[j])
    for i, month in enumerate(months):
        a = monthly_a[i] if i < len(monthly_a) else 0
        b = monthly_b[i] if i < len(monthly_b) else 0
        c = monthly_c[i] if i < len(monthly_c) else a - b
        ec.rows[i + 1].cells[0].text = month
        ec.rows[i + 1].cells[1].text = f"{a:,.2f}"
        ec.rows[i + 1].cells[2].text = f"{c:,.2f}"
        ec.rows[i + 1].cells[3].text = f"{b:,.2f}"
    for j, v in enumerate(["TOTAL", f"{total_a:,.2f}", f"{total_c:,.2f}", f"{total_b:,.2f}"]):
        ec.rows[n + 1].cells[j].text = v
        _bold_cell(ec.rows[n + 1].cells[j])
    doc.add_paragraph()

    _heading(doc, "4.2 Monthly Electricity Cost Estimate (RM)", level=2)
    cost_rows = []
    total_cost = 0.0
    for i, month in enumerate(months):
        a = monthly_a[i] if i < len(monthly_a) else 0
        cost = a * tariff
        total_cost += cost
        cost_rows.append((month, f"{cost:,.2f}"))
    ct = doc.add_table(rows=n + 2, cols=2)
    ct.style = "Table Grid"
    for j, h in enumerate(["Month", "Est. Consumption Charges (RM)"]):
        ct.rows[0].cells[j].text = h
        _bold_cell(ct.rows[0].cells[j])
    for i, (m, c) in enumerate(cost_rows):
        ct.rows[i + 1].cells[0].text = m
        ct.rows[i + 1].cells[1].text = c
    ct.rows[n + 1].cells[0].text = "TOTAL"
    ct.rows[n + 1].cells[1].text = f"{total_cost:,.2f}"
    _bold_cell(ct.rows[n + 1].cells[0])
    _bold_cell(ct.rows[n + 1].cells[1])
    doc.add_page_break()

    # ── 5. GFA ────────────────────────────────────────────────────────────────
    _heading(doc, "5. Gross Floor Area (GFA) Calculation")
    _para(doc, (
        f"The Gross Floor Area (GFA) of {building} was determined in accordance with "
        f"Suruhanjaya Tenaga guidelines, encompassing all conditioned floor areas within "
        f"the building boundary."
    ))
    _para(doc, f"Declared GFA:  {gfa:,.2f} m²", bold=True)
    doc.add_page_break()

    # ── 6. BEI Calculation ────────────────────────────────────────────────────
    _heading(doc, "6. Building Energy Intensity (BEI) Calculation")
    for blk in narratives.get("bei_analysis", "").split("\n\n"):
        if blk.strip():
            _para(doc, blk.strip())
    doc.add_paragraph()
    _para(doc, "BEI  =  Total Supply to Building (A)  ÷  Gross Floor Area", bold=True)
    _para(doc, f"     =  {total_a:,.2f} kWh  ÷  {gfa:,.2f} m²")
    _para(doc, f"     =  {bei} kWh/m²/year  =  {bei_gj} GJ/m²/year  →  {star}", bold=True)
    _para(doc, "")
    _para(doc, f"Net Landlord Load (C = A − B)  =  {total_c:,.2f} kWh", bold=True)
    _para(doc, f"Tenant Load (B)  =  {total_b:,.2f} kWh")
    doc.add_paragraph()
    for blk in narratives.get("landlord_tenant_analysis", "").split("\n\n"):
        if blk.strip():
            _para(doc, blk.strip())
    doc.add_page_break()

    # ── 7. Conclusion & Recommendations ──────────────────────────────────────
    _heading(doc, "7. Conclusion and Recommendations")
    recs = narratives.get("conclusions_recommendations", "")
    for line in recs.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith(("•", "-", "*")):
            doc.add_paragraph(line.lstrip("•-* ").strip(), style="List Bullet")
        else:
            _para(doc, line)
    doc.add_page_break()

    # ── 8. Verification ───────────────────────────────────────────────────────
    _heading(doc, "8. Verification")
    _para(doc, (
        "This report has been prepared, reviewed, and submitted in accordance with "
        "Suruhanjaya Tenaga (ST) requirements for Building Energy Intensity reporting."
    ))
    doc.add_paragraph()
    _table2(doc, [
        ("Prepared By",  prep_name),
        ("Position",     prep_pos),
        ("Company",      "Atech Energy Sdn Bhd"),
        ("Date",         sub_date),
        ("Signature",    ""),
    ])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
